from docx import Document
from constants.constants import Constants

class ResumeService:

    def __init__(self):

        constants = Constants()
        self.overview_section = constants.overview_section
        self.technologies_section = constants.technologies_section

    def match_skills(self, job_description, resume_skills):

        matched_skills = [skill for skill in resume_skills if skill.lower() in job_description.lower()]
        return matched_skills


    def update_resume(self, doc_path, tailored_section, matched_skills):

        tailored_document = Document(doc_path)

        for paragraph in tailored_document.paragraphs:
            if self.overview_section in paragraph.text:
                paragraph.clear()
                paragraph.add_run(tailored_section)

        for paragraph in tailored_document.paragraphs:
            if self.technologies_section in paragraph.text:
                paragraph.clear()
                paragraph.add_run("\n Matched Skills: " + ", ".join(matched_skills))

        return tailored_document