from dotenv import load_dotenv
import os
from openai import OpenAI
from docx import Document
from constants.constants import Constants

class OpenAiService:

    def __init__(self):

        constants = Constants()
        load_dotenv('../.env')
        self.client = self._get_client()
        self.model = constants.model
        self.personal_name = constants.personal_name
        self.resume_skills = constants.skills


    def match_skills(self, job_description, resume_skills):

        matched_skills = [skill for skill in resume_skills if skill.lower() in job_description.lower()]
        return matched_skills


    def generate_resume_bullet_points(self, job_description, matched_skills):

        prompt = (
            f"Based on the following job description:\n{job_description}\n\n"
            f"and these skills: {', '.join(matched_skills)}, "
            f"write 3 concise and professional resume bullet points tailored to this job."
        )

        return self._get_content(prompt)


    def generate_cover_letter(self, job_description, matched_skills):

        prompt = (
            f"Based on the following job description:\n{job_description}\n\n"
            f"and these skills: {', '.join(matched_skills)}, "
            f"write a short paragraph for a cover letter tailored to this job. Sign it with {self.personal_name}"
        )

        return self._get_content(prompt)


    def generate_resume(self, job_description, matching_skills):

        prompt = (
            f"Based on the following job description:\n{job_description}\n\n"
            f"and these skills: {', '.join(matching_skills)}, "
            f"write a professional resume section highlighting relevant achievements and experience."
            f"do not include education, only have overview and key achievements. do not include any job specifically"
        )

        return self._get_content(prompt)


    def get_list_of_skills_from_job_description(self, job_description):

        prompt = (
            f"Based on the following job description:\n {job_description}\n\n"
            f"Parse as many skills as possible for a resume as a comma seperated list"
        )
        return self._get_content(prompt)


    def update_resume(self, doc_path, tailored_section, matched_skills):
        doc = Document(doc_path)
        for paragraph in doc.paragraphs:
            if "Overview" in paragraph.text:
                paragraph.clear()
                paragraph.add_run("Overview: \n " + tailored_section)
                #paragraph.text = "Overview: \n" + tailored_section

        for paragraph in doc.paragraphs:
            if "Technologies" in paragraph.text:
                paragraph.clear()
                paragraph.add_run("\n Matched Skills: " + ", ".join(matched_skills))
                #paragraph.text += "\nMatched Skills: " + ", ".join(matched_skills)

        return doc

    def get_resume_skills(self):
        return self.resume_skills

    def _get_content(self, prompt):
        response = self.client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model=self.model,
        )
        choices = response.choices
        return choices[0].message.content

    def _get_client(self):

        OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
        if not OPENAI_API_KEY:
            raise Exception("Missing OPENAI_API_KEY environment variable")
        client = OpenAI(
            api_key=OPENAI_API_KEY
        )
        return client

